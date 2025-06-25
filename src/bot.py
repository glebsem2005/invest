import html
import logging
import re
import traceback
import json
import os
import smtplib
from abc import ABC, abstractmethod
from typing import Any, Dict, Tuple
from docx import Document
from docx.shared import Inches
import tempfile
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

import aiogram.utils.exceptions
from aiogram import Bot, Dispatcher, executor, types
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup

from access_middleware import AccessMiddleware
from chat_context import ChatContextManager
from config import Config
from file_processor import FileProcessor
from keyboards_builder import Button, DynamicKeyboard, Keyboard
from logger import Logger
from models_api import ExcelFileManager, ExcelSearchStrategy, ModelAPI
from prompts import DEFAULT_PROMPTS_DIR, Models, SystemPrompt, SystemPrompts, Topics
from sql_auth import init_auth_system, check_user_authorized, get_user_email

Logger()
logger = logging.getLogger('bot')
config = Config()

# =============================================================================
# СОСТОЯНИЯ
# =============================================================================

class UserStates(StatesGroup):
    ACCESS = State()
    CHOOSING_TOPIC = State()
    CHOOSING_MODEL = State()
    ENTERING_PROMPT = State()
    ATTACHING_FILE = State()
    UPLOADING_FILE = State()
    ASKING_CONTINUE = State()
    CONTINUE_DIALOG = State()
    ATTACHING_FILE_CONTINUE = State()
    UPLOADING_FILE_CONTINUE = State()
    
    # Инвестиционный анализ
    INVESTMENT_ACTIONS = State()
    INVESTMENT_QA = State()
    INVESTMENT_REPORT_OPTIONS = State()
    CHOOSING_FINAL_ACTION = State()

class AdminStates(StatesGroup):
    CHOOSING_PROMPT = State()
    CHOOSING_PROMPT_TYPE = State()
    UPLOADING_SYSTEM_PROMPT = State()
    UPLOADING_DETAIL_PROMPT = State()
    UPLOADING_PROMPT = State()
    NEW_PROMPT_NAME = State()
    NEW_PROMPT_DISPLAY = State()
    NEW_PROMPT_UPLOAD = State()
    NEW_PROMPT_UPLOAD_DETAIL = State()
    UPLOADING_SCOUTING_FILE = State()
    CHOOSING_AI_MODEL = State()

# =============================================================================
# КЛАВИАТУРЫ
# =============================================================================

class TopicKeyboard(DynamicKeyboard):
    @classmethod
    def get_buttons(cls) -> Tuple[Button, ...]:
        buttons = []
        for topic_name, topic in Topics.__members__.items():
            buttons.append(Button(text=topic.value, callback=f'topic_{topic_name}'))
        return tuple(buttons)

class FileAttachKeyboard(DynamicKeyboard):
    @classmethod
    def get_buttons(cls) -> Tuple[Button, ...]:
        return (
            Button(text='Да, прикрепить файл', callback='attach_file'),
            Button(text='Нет, продолжить без файла', callback='no_file'),
        )

class ContinueKeyboard(Keyboard):
    _buttons = (
        Button('Задать вопрос', 'continue_yes'),
        Button('Завершить чат', 'continue_no'),
    )

class UnauthorizedKeyboard(Keyboard):
    _buttons = (Button('🔐 Пройти авторизацию', 'go_to_main_bot'),)
    
    @classmethod
    def get_markup(cls):
        markup = types.InlineKeyboardMarkup(row_width=1)
        markup.add(types.InlineKeyboardButton(
            text="🔐 Пройти авторизацию",
            url="https://t.me/sberallaibot"
        ))
        return markup

class InvestmentActionsKeyboard(Keyboard):
    _buttons = (
        Button('Повторная генерация', 'investment_regenerate'),
        Button('Задать вопрос', 'investment_ask_question'),
        Button('Получить отчет', 'investment_get_report'),
    )

class InvestmentReportKeyboard(Keyboard):
    _buttons = (
        Button('Скачать отчет', 'investment_download'),
        Button('Выслать на почту', 'investment_email'),
        Button('← Назад', 'investment_back_to_actions'),
    )

class FinalActionsKeyboard(Keyboard):
    _buttons = (
        Button('🏢 Новая компания', 'new_company_analysis'),
        Button('← Вернуться к основному боту', 'return_to_main_bot'),
    )

class AdminPromptKeyboard(DynamicKeyboard):
    @classmethod
    def get_buttons(cls) -> Tuple[Button, ...]:
        buttons = []
        for topic_name, topic in Topics.__members__.items():
            buttons.append(Button(text=topic.value, callback=f'prompt_{topic_name}'))
        return tuple(buttons)

class PromptTypeKeyboard(Keyboard):
    _buttons = (
        Button('Системный промпт', 'prompt_type_system'),
        Button('Детализированный промпт', 'prompt_type_detail'),
        Button('Оба промпта', 'prompt_type_both'),
    )

class AdminAIModelKeyboard(DynamicKeyboard):
    @classmethod
    def get_buttons(cls) -> Tuple[Button, ...]:
        buttons = []
        for model_name, model in Models.__members__.items():
            buttons.append(Button(text=f"Модель: {model_name.upper()}", callback=f'ai_model_{model_name}'))
        return tuple(buttons)

# =============================================================================
# EMAIL SENDER
# =============================================================================

class EmailSender:
    """Класс для отправки email с отчетами."""
    
    def __init__(self):
        self.smtp_server = config.SMTP_SERVER
        self.smtp_port = config.SMTP_PORT  
        self.email_user = config.EMAIL_USER
        self.email_password = config.EMAIL_PASSWORD
        self.sender_name = config.SENDER_NAME
        
        if not self.email_user or not self.email_password:
            logger.warning("Email credentials not configured. Email sending will not work.")
    
    def _sanitize_filename(self, filename: str) -> str:
        """Очищает имя файла от недопустимых символов."""
        if not filename or not filename.strip():
            return 'unknown_company'
        
        sanitized = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', str(filename))
        sanitized = sanitized.strip().replace(' ', '_')
        
        if len(sanitized) > 50:
            sanitized = sanitized[:50]
        sanitized = sanitized.rstrip('.')
        
        if not sanitized:
            return 'unknown_company'
        
        return sanitized

    async def send_report(self, recipient_email: str, company_name: str, report_file_path: str, filename: str = None) -> bool:
        """Отправляет отчет на указанный email."""
        if not self.email_user or not self.email_password:
            logger.error("Email credentials not configured")
            return False
        
        try:
            msg = MIMEMultipart()
            msg['From'] = f"{self.sender_name} <{self.email_user}>"
            msg['To'] = recipient_email
            msg['Subject'] = f"Инвестиционный анализ: {company_name}"
        
            body = f"""Здравствуйте!

Высылаем вам результаты инвестиционного анализа компании "{company_name}".

Отчет содержит:
- Анализ рынка и конкурентов
- Оценку инвестиционной привлекательности  
- Рекомендации по взаимодействию

С уважением,
Команда аналитиков"""
        
            msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
            if os.path.exists(report_file_path):
                with open(report_file_path, "rb") as attachment:
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(attachment.read())
                    encoders.encode_base64(part)
            
                safe_filename = filename if filename else f"investment_analysis_{self._sanitize_filename(company_name)}.docx"
                if not safe_filename:
                    safe_filename = "investment_analysis_report.docx"
            
                part.add_header('Content-Disposition', f'attachment; filename="{safe_filename}"')
                msg.attach(part)
                logger.info(f"Attached file with name: {safe_filename}")
            else:
                logger.error(f"Report file not found: {report_file_path}")
                return False
        
            server = smtplib.SMTP(self.smtp_server, self.smtp_port)
            server.starttls()
            server.login(self.email_user, self.email_password)
            text = msg.as_string()
            server.sendmail(self.email_user, recipient_email, text)
            server.quit()
        
            logger.info(f"Email successfully sent to {recipient_email}")
            return True
        
        except Exception as e:
            logger.error(f"Failed to send email to {recipient_email}: {e}")
            return False

# =============================================================================
# INVESTMENT ANALYSIS PROCESSOR
# =============================================================================

class InvestmentAnalysisProcessor:
    """Класс для обработки анализа инвестиционной привлекательности."""
    
    def __init__(self):
        self.analysis_prompt = """
Найди название компании в тексте и определи типы анализа.

ГЛАВНАЯ ЗАДАЧА: точно определить название компании.

Текст: "{user_text}"

Инструкции:
1. Найди название компании или бренда в тексте
2. Если названия нет, но есть описание ("фудтех стартап"), напиши "неизвестная_компания"
3. Определи нужные анализы:
   - market: 1 если нужен рыночный анализ (рынок, финансы, позиция)
   - rivals: 1 если нужен анализ конкурентов
   - synergy: 1 если нужен анализ синергии
   - Если тип анализа не указан, ставь все в 1

Ответ только JSON: {{"name": "название", "market": 1, "rivals": 1, "synergy": 1}}
"""
        
        self.executive_summary_prompt = """
1. РОЛЬ

• Ты — инвестиционный профессионал с 50-летним опытом в корпоративном развитии и M&A, с глубоким опытом разработки экосистемных стратегий для Big Tech (Apple, Amazon, Yandex, Baidu, Tencent и пр.) и глубокими знаниями экономики на уровне нобелевских лауреатов по экономике.
• В твои обязанности входит поиск направлений развития экосистемы Сбера через M&A и стратегические партнерства.

2. ЗАДАЧА И ФОРМАТ ОТВЕТА

Проанализируй целесообразность партнерства Сбера с компанией [Название компании]. 

Ответ должен быть в формате executive summary для высшего руководства банка – простой текст объемом не более 250 слов с разбивкой на несколько абзацев. Ключевые выводы должны быть подтверждены числовыми данными, рыночной аналитикой, денежной оценкой.

Будь конкретным и структурированным в ответе.
"""

    def _get_ai_model(self):
        """Получает настроенную модель AI из конфигурации."""
        try:
            system_prompts = SystemPrompts()
            ai_model_config = system_prompts.get_prompt_if_exists('AI_MODEL_CONFIG')
            if ai_model_config:
                model_name = ai_model_config.strip().lower()
                if hasattr(Models, model_name):
                    return Models[model_name].value()
            return Models.chatgpt.value()
        except:
            return Models.chatgpt.value()

    async def parse_user_request(self, user_text: str) -> Dict[str, Any]:
        """Парсит запрос пользователя и определяет параметры анализа."""
        try:
            model_api = ModelAPI(Models.chatgpt.value())
            messages = [
                {"role": "system", "content": "Ты помощник для извлечения названий компаний из текста. Отвечай только валидным JSON без лишнего текста."},
                {"role": "user", "content": self.analysis_prompt.format(user_text=user_text)}
            ]
            
            response = await model_api.get_response(messages)
            logger.info(f"Raw response from GPT: '{response}'")
            
            # Попытка парсинга JSON
            try:
                result = json.loads(response.strip())
                if "name" in result and result["name"].strip():
                    return {
                        "name": result.get("name", "неизвестная_компания"),
                        "market": result.get("market", 1),
                        "rivals": result.get("rivals", 1), 
                        "synergy": result.get("synergy", 1)
                    }
            except json.JSONDecodeError:
                pass
            
        except Exception as e:
            logger.error(f"Error parsing user request: {e}")
        
        # Fallback
        fallback_result = {"name": "неизвестная_компания", "market": 1, "rivals": 1, "synergy": 1}
        logger.info(f"Using fallback result: {fallback_result}")
        return fallback_result

    async def run_analysis(self, analysis_params: Dict[str, Any], file_content: str = "") -> Dict[str, str]:
        """Запускает анализ согласно параметрам."""
        results = {}
        system_prompts = SystemPrompts()
        model_api = ModelAPI(self._get_ai_model())
        company_name = analysis_params.get("name", "unknown_company")
        
        additional_context = f"\n\nДополнительная информация из файла:\n{file_content}" if file_content else ""
        analysis_context = ""
        
        # Рыночный анализ
        if analysis_params.get("market", 0):
            try:
                market_prompt_raw = system_prompts.get_prompt(SystemPrompt.INVESTMENT_MARKET)
                parsed_prompt = self._parse_classical_prompt(market_prompt_raw)
                
                system_content = parsed_prompt["role"] + "\n\nОТВЕТ ДОЛЖЕН БЫТЬ СТРОГО НЕ БОЛЕЕ 300 СЛОВ."
                user_content = parsed_prompt["prompt"].replace("[название компании]", company_name) + additional_context
                
                messages = [
                    {"role": "system", "content": system_content},
                    {"role": "user", "content": user_content}
                ]
                
                results["market"] = await model_api.get_response(messages)
                analysis_context += f"\n\nРезультат рыночного анализа компании {company_name}:\n{results['market']}"
                logger.info("Market analysis completed")
            except Exception as e:
                logger.error(f"Error in market analysis: {e}")
                results["market"] = f"Ошибка при анализе рынка: {str(e)}"
        
        # Анализ конкурентов
        if analysis_params.get("rivals", 0):
            try:
                rivals_prompt_raw = system_prompts.get_prompt(SystemPrompt.INVESTMENT_RIVALS)
                parsed_prompt = self._parse_classical_prompt(rivals_prompt_raw)
                
                system_content = parsed_prompt["role"] + "\n\nОТВЕТ ДОЛЖЕН БЫТЬ СТРОГО НЕ БОЛЕЕ 300 СЛОВ."
                user_content = parsed_prompt["prompt"].replace("[название компании]", company_name) + additional_context + analysis_context
                
                messages = [
                    {"role": "system", "content": system_content},
                    {"role": "user", "content": user_content}
                ]
                
                results["rivals"] = await model_api.get_response(messages)
                analysis_context += f"\n\nРезультат анализа конкурентов компании {company_name}:\n{results['rivals']}"
                logger.info("Rivals analysis completed")
            except Exception as e:
                logger.error(f"Error in rivals analysis: {e}")
                results["rivals"] = f"Ошибка при анализе конкурентов: {str(e)}"
        
        # Анализ синергии
        if analysis_params.get("synergy", 0):
            try:
                synergy_prompt_raw = system_prompts.get_prompt(SystemPrompt.INVESTMENT_SYNERGY)
                
                if isinstance(synergy_prompt_raw, dict):
                    system_content = synergy_prompt_raw.get("role", "") + "\n\nОТВЕТ ДОЛЖЕН БЫТЬ СТРОГО НЕ БОЛЕЕ 300 СЛОВ."
                    user_content = synergy_prompt_raw.get("prompt", "").replace("[название компании]", company_name) + additional_context + analysis_context
                else:
                    parsed_prompt = self._parse_classical_prompt(synergy_prompt_raw)
                    system_content = parsed_prompt["role"] + "\n\nОТВЕТ ДОЛЖЕН БЫТЬ СТРОГО НЕ БОЛЕЕ 300 СЛОВ."
                    user_content = parsed_prompt["prompt"].replace("[название компании]", company_name) + additional_context + analysis_context
                
                messages = [
                    {"role": "system", "content": system_content},
                    {"role": "user", "content": user_content}
                ]
                
                results["synergy"] = await model_api.get_response(messages)
                logger.info("Synergy analysis completed")
            except Exception as e:
                logger.error(f"Error in synergy analysis: {e}")
                results["synergy"] = f"Ошибка при анализе синергии: {str(e)}"
        
        return results
    
    def _parse_classical_prompt(self, prompt_text: str) -> Dict[str, str]:
        """Парсит классический промпт формата 'РОЛЬ. ... КОНТЕКСТ. ...'"""
        try:
            role_match = re.search(r'РОЛЬ\.\s*(.*?)(?=\n\s*КОНТЕКСТ\.|\n\s*[А-ЯЁ]+\.|\Z)', prompt_text, re.DOTALL | re.IGNORECASE)
            context_match = re.search(r'КОНТЕКСТ\.\s*(.*?)(?=\n\s*[А-ЯЁ]+\.|\Z)', prompt_text, re.DOTALL | re.IGNORECASE)
            
            role = role_match.group(1).strip() if role_match else "Ты профессиональный аналитик."
            context = context_match.group(1).strip() if context_match else prompt_text.strip()
            
            return {"role": role, "prompt": context}
        except Exception as e:
            logger.warning(f"Error parsing classical prompt: {e}")
            return {"role": "Ты профессиональный аналитик.", "prompt": prompt_text}

    def create_docx_report(self, company_name: str, analysis_results: Dict[str, str]) -> str:
        """Создает DOCX отчет с результатами анализа."""
        try:
            doc = Document()
            doc.add_heading(f'Анализ инвестиционной привлекательности: {company_name}', 0)
            
            if "market" in analysis_results:
                doc.add_heading('Рыночный анализ', level=1)
                doc.add_paragraph(analysis_results["market"])
                doc.add_page_break()
            
            if "rivals" in analysis_results:
                doc.add_heading('Анализ конкурентов', level=1)
                doc.add_paragraph(analysis_results["rivals"])
                doc.add_page_break()
            
            if "synergy" in analysis_results:
                doc.add_heading('Анализ синергии', level=1)
                doc.add_paragraph(analysis_results["synergy"])
            
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            temp_file.close()
            
            logger.info(f"DOCX report created: {temp_file.name}")
            return temp_file.name
            
        except Exception as e:
            logger.error(f"Error creating DOCX report: {e}")
            raise

    async def generate_executive_summary(self, docx_file_path: str) -> str:
        """Генерирует executive summary на основе DOCX файла."""
        try:
            doc = Document(docx_file_path)
            doc_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            model_api = ModelAPI(Models.chatgpt.value())
            messages = [
                {"role": "system", "content": self.executive_summary_prompt},
                {"role": "user", "content": f"Содержимое анализа:\n\n{doc_content}"}
            ]
            
            executive_summary = await model_api.get_response(messages)
            logger.info("Executive summary generated")
            return executive_summary
            
        except Exception as e:
            logger.error(f"Error generating executive summary: {e}")
            return "Ошибка при генерации executive summary"

    async def create_final_report_with_qa(self, company_name: str, analysis_results: Dict[str, str], qa_history: list) -> str:
        """Создает финальный отчет с интегрированными Q&A."""
        try:
            doc = Document()
            
            title = doc.add_heading(f'Инвестиционный анализ: {company_name}', 0)
            title.alignment = 1
        
            from datetime import datetime
            date_paragraph = doc.add_paragraph(f'Дата создания отчета: {datetime.now().strftime("%d.%m.%Y")}')
            date_paragraph.alignment = 1
            doc.add_page_break()
        
            if "market" in analysis_results:
                doc.add_heading('1. Рыночный анализ', level=1)
                doc.add_paragraph(analysis_results["market"])
                doc.add_page_break()
        
            if "rivals" in analysis_results:
                doc.add_heading('2. Анализ конкурентов', level=1)
                doc.add_paragraph(analysis_results["rivals"])
                doc.add_page_break()
        
            if "synergy" in analysis_results:
                doc.add_heading('3. Анализ синергии', level=1)
                doc.add_paragraph(analysis_results["synergy"])
                if qa_history:
                    doc.add_page_break()
        
            if qa_history:
                doc.add_heading('4. Дополнительные вопросы и ответы', level=1)
                for i, qa in enumerate(qa_history, 1):
                    question_para = doc.add_paragraph()
                    question_run = question_para.add_run(f"Вопрос {i}: ")
                    question_run.bold = True
                    question_para.add_run(qa['question'])
                
                    answer_para = doc.add_paragraph()
                    answer_run = answer_para.add_run("Ответ: ")
                    answer_run.bold = True
                    answer_para.add_run(qa['answer'])
                    
                    if i < len(qa_history):
                        doc.add_paragraph()
        
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='_final.docx')
            doc.save(temp_file.name)
            temp_file.close()
        
            logger.info(f"Final report with Q&A created: {temp_file.name}")
            return temp_file.name
        
        except Exception as e:
            logger.error(f"Error creating final report with Q&A: {e}")
            return self.create_docx_report(company_name, analysis_results)

    def _sanitize_filename(self, filename: str) -> str:
        """Очищает имя файла от недопустимых символов."""
        if not filename or not filename.strip():
            return 'unknown_company'
        
        sanitized = re.sub(r'[<>:"/\\|?*\x00-\x1f]', '_', str(filename))
        sanitized = sanitized.strip().replace(' ', '_')
        
        if len(sanitized) > 50:
            sanitized = sanitized[:50]
        sanitized = sanitized.rstrip('.')
        
        return sanitized if sanitized else 'unknown_company'

# =============================================================================
# BASE SCENARIO
# =============================================================================

class BaseScenario(ABC):
    """Базовый класс для сценариев с общей логикой работы с запросами, файлами и ошибками."""

    def __init__(self, bot: Bot) -> None:
        self.bot = bot

    @abstractmethod
    async def process(self, *args, **kwargs) -> Any:
        pass

    @abstractmethod
    def register(self, dp: Dispatcher) -> None:
        pass

    async def process_investment_analysis(self, message, state, file_content=''):
        """Обработка запроса для анализа инвестиционной привлекательности."""
        user_id = message.chat.id
        user_data = await state.get_data()
        user_query = user_data.get('user_query', '')

        await self.delete_message_by_id(user_id, user_data.get('processing_msg_id'))

        if not user_query and not file_content:
            await message.answer('Необходимо ввести запрос или прикрепить файл. Пожалуйста, начните заново с команды /start')
            return

        try:
            progress_msg = await message.answer('🔍 Анализирую запрос и определяю параметры анализа...')
            
            processor = InvestmentAnalysisProcessor()
            analysis_params = await processor.parse_user_request(user_query)
            company_name = analysis_params.get("name", "unknown_company")
            
            await progress_msg.edit_text(f'📊 Запускаю анализ для компании: {company_name}...')
            
            analysis_results = await processor.run_analysis(analysis_params, file_content)
            
            await progress_msg.edit_text('📄 Создаю отчет...')
            docx_file_path = processor.create_docx_report(company_name, analysis_results)
            
            await progress_msg.edit_text('📝 Генерирую executive summary...')
            executive_summary = await processor.generate_executive_summary(docx_file_path)
            
            await state.update_data(
                analysis_params=analysis_params,
                analysis_results=analysis_results,
                docx_file_path=docx_file_path,
                executive_summary=executive_summary,
                company_name=company_name,
                qa_history=[]
            )
            
            await self.send_markdown_response(message, executive_summary)
            await progress_msg.delete()
            
            await message.answer('Что бы вы хотели сделать дальше?', reply_markup=InvestmentActionsKeyboard())
            await UserStates.INVESTMENT_ACTIONS.set()
            
        except Exception as e:
            await self.handle_error(message, e, "investment_analysis")

    async def send_markdown_response(self, message, response):
        """Отправляет ответ с markdown форматированием."""
        escaped_response = self._escape_markdown(response)
        max_length = 4000
        for i in range(0, len(escaped_response), max_length):
            part = escaped_response[i : i + max_length]
            await message.answer(part, parse_mode='MarkdownV2')

    async def delete_message_by_id(self, user_id, message_id):
        """Удаляет сообщение по ID."""
        if message_id:
            try:
                await self.bot.delete_message(chat_id=user_id, message_id=message_id)
            except Exception:
                pass

    async def handle_error(self, message, e, model_name):
        """Обрабатывает ошибки."""
        logger.error(f'Ошибка {model_name}: {e}', exc_info=True)

        token_limit = self._parse_token_limit_error(str(e))
        if token_limit:
            await message.answer(f'⚠️ Вы превысили лимит токенов для модели.\nМаксимум: {token_limit} токенов.\nПожалуйста, уменьшите запрос и попробуйте снова.')
        else:
            await message.answer(
                'Извините, мой маленький компьютер перегружен. Поступает слишком много запросов. Пожалуйста, подождите несколько секунд или попробуйте ещё раз.\n'
                'Если проблема не исчезнет, обратитесь к администратору.'
            )

        await self.bot.send_message(chat_id=config.OWNER_ID, text=f'Ошибка при запросе к {model_name} от пользователя {message.chat.id}:\n{e}')

    def _escape_markdown(self, text: str) -> str:
        """Экранирует специальные символы для MarkdownV2."""
        try:
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            result = ''
            i = 0
            while i < len(text):
                if i + 1 < len(text) and text[i : i + 2] == '**':
                    end_pos = text.find('**', i + 2)
                    if end_pos != -1:
                        bold_content = text[i + 2 : end_pos]
                        escaped_content = ''
                        for char in bold_content:
                            if char in '_[]()~`>#+-=|{}.!':
                                escaped_content += f'\\{char}'
                            elif char == '*':
                                escaped_content += '\\*'
                            else:
                                escaped_content += char
                        result += f'*{escaped_content}*'
                        i = end_pos + 2
                        continue
                if text[i] in '_*[]()~`>#+-=|{}.!':
                    result += f'\\{text[i]}'
                else:
                    result += text[i]
                i += 1
            return result
        except Exception:
            return text

    def _parse_token_limit_error(self, error_text: str) -> int:
        """Парсит ошибку лимита токенов."""
        match = re.search(r'Limit (\d+), Requested (\d+)', error_text)
        if match:
            return int(match.group(1))
        return None

# =============================================================================
# INVESTMENT HANDLERS
# =============================================================================

class InvestmentActionsHandler(BaseScenario):
    """Обработка действий после получения executive summary."""

    async def process(self, callback_query, state, **kwargs) -> None:
        user_id = callback_query.from_user.id
        action = callback_query.data
        user_data = await state.get_data()
        
        logger.info(f"Investment actions handler: user {user_id}, action {action}")
        await callback_query.answer()

        if action == 'investment_regenerate':
            await callback_query.message.delete()
            progress_msg = await callback_query.message.answer('🔄 Запускаю повторную генерацию анализа...')
            
            try:
                processor = InvestmentAnalysisProcessor()
                analysis_params = user_data.get('analysis_params')
                company_name = user_data.get('company_name')
                
                analysis_results = await processor.run_analysis(analysis_params)
                docx_file_path = processor.create_docx_report(company_name, analysis_results)
                executive_summary = await processor.generate_executive_summary(docx_file_path)
                
                await state.update_data(
                    analysis_results=analysis_results,
                    docx_file_path=docx_file_path,
                    executive_summary=executive_summary,
                    qa_history=[]
                )
                
                await progress_msg.delete()
                await self.send_markdown_response(callback_query.message, executive_summary)
                await callback_query.message.answer('Что бы вы хотели сделать дальше?', reply_markup=InvestmentActionsKeyboard())
                
            except Exception as e:
                await progress_msg.delete()
                await self.handle_error(callback_query.message, e, "regeneration")
            
        elif action == 'investment_ask_question':
            await callback_query.message.delete()
            await callback_query.message.answer('❓ Задайте ваш вопрос по анализу компании. Все вопросы и ответы будут добавлены в итоговый отчет.')
            await UserStates.INVESTMENT_QA.set()
            
        elif action == 'investment_get_report':
            await callback_query.message.edit_text('Выберите способ получения отчета:', reply_markup=InvestmentReportKeyboard())
            await UserStates.INVESTMENT_REPORT_OPTIONS.set()

    def register(self, dp: Dispatcher) -> None:
        dp.register_callback_query_handler(
            self.process,
            lambda c: c.data in ['investment_regenerate', 'investment_ask_question', 'investment_get_report'],
            state=UserStates.INVESTMENT_ACTIONS,
        )

class InvestmentQAHandler(BaseScenario):
    """Обработка вопросов-ответов в режиме инвестиционного анализа."""

    async def process(self, message, state, **kwargs) -> None:
        user_id = message.from_user.id
        user_data = await state.get_data()
        user_question = message.text
        
        company_name = user_data.get('company_name', 'неизвестная_компания')
        qa_history = user_data.get('qa_history', [])

        try:
            model_api = ModelAPI(Models.chatgpt.value())
            messages = [
                {"role": "system", "content": f"Ты эксперт по инвестиционному анализу. Ответь на вопрос по компании {company_name}. Будь конкретным и профессиональным."},
                {"role": "user", "content": user_question}
            ]
            
            await self.bot.send_chat_action(chat_id=user_id, action='typing')
            response = await model_api.get_response(messages)
            
            qa_history.append({"question": user_question, "answer": response})
            await state.update_data(qa_history=qa_history)
            
            await self.send_markdown_response(message, response)
            
            await message.answer(
                'Хотите задать еще вопрос или вернуться к выбору действий?',
                reply_markup=types.InlineKeyboardMarkup().add(
                    types.InlineKeyboardButton('← Вернуться к действиям', callback_data='back_to_investment_actions')
                )
            )
            
        except Exception as e:
            await self.handle_error(message, e, "investment_qa")

    def register(self, dp: Dispatcher) -> None:
        dp.register_message_handler(self.process, content_types=['text'], state=UserStates.INVESTMENT_QA)

class BackToInvestmentActionsHandler(BaseScenario):
    """Возврат к выбору действий инвестиционного анализа."""

    async def process(self, callback_query, state, **kwargs) -> None:
        await callback_query.answer()
        await callback_query.message.edit_text('Что бы вы хотели сделать дальше?', reply_markup=InvestmentActionsKeyboard())
        await UserStates.INVESTMENT_ACTIONS.set()

    def register(self, dp: Dispatcher) -> None:
        dp.register_callback_query_handler(
            self.process,
            lambda c: c.data == 'back_to_investment_actions',
            state='*',
        )

class InvestmentReportHandler(BaseScenario):
    """Обработка получения отчета."""

    def __init__(self, bot):
        super().__init__(bot)
        self.email_sender = EmailSender()

    async def process(self, callback_query, state, **kwargs) -> None:
        user_id = callback_query.from_user.id
        action = callback_query.data
        user_data = await state.get_data()

        await callback_query.answer()

        if action == 'investment_download':
            await self._download_report(callback_query, state, user_data)
        elif action == 'investment_email':
            await self._send_email_report(callback_query, state, user_data)
        elif action == 'investment_back_to_actions':
            await callback_query.message.edit_text('Что бы вы хотели сделать дальше?', reply_markup=InvestmentActionsKeyboard())
            await UserStates.INVESTMENT_ACTIONS.set()

    async def _download_report(self, callback_query, state, user_data):
        """Генерирует и отправляет отчет для скачивания."""
        try:
            await callback_query.message.edit_text('Генерирую финальный отчет...')
            
            processor = InvestmentAnalysisProcessor()
            company_name = user_data.get('company_name', 'unknown_company')
            analysis_results = user_data.get('analysis_results')
            qa_history = user_data.get('qa_history', [])
            
            final_report_path = await processor.create_final_report_with_qa(company_name, analysis_results, qa_history)
            
            safe_company_name = processor._sanitize_filename(company_name)
            report_filename = f'investment_analysis_{safe_company_name}_final.docx'
            
            with open(final_report_path, 'rb') as doc_file:
                await callback_query.message.answer_document(
                    document=types.InputFile(doc_file, filename=report_filename),
                    caption=f'Финальный отчет c инвестиционным анализом: {company_name}'
                )
            
            os.unlink(final_report_path)
            
            await callback_query.message.answer('Отчет готов! Что бы вы хотели сделать дальше?', reply_markup=FinalActionsKeyboard())
            await UserStates.CHOOSING_FINAL_ACTION.set()
            
        except Exception as e:
            await self.handle_error(callback_query.message, e, "report_generation")

    async def _send_email_report(self, callback_query, state, user_data):
        """Автоматически отправляет отчет на email из БД."""
        user_id = callback_query.from_user.id
        
        if not self.email_sender.email_user or not self.email_sender.email_password:
            await callback_query.message.edit_text('❌ Отправка на email временно недоступна. Воспользуйтесь скачиванием отчета.', reply_markup=InvestmentReportKeyboard())
            return
        
        try:
            await callback_query.message.edit_text('🔍 Получаю ваш email из базы данных...')
            
            user_email = await get_user_email(user_id)
            
            if not user_email:
                await callback_query.message.edit_text('❌ Не удалось получить ваш email из базы данных.\nВоспользуйтесь скачиванием отчета.', reply_markup=InvestmentReportKeyboard())
                return
            
            await callback_query.message.edit_text(f'📧 Генерирую и отправляю отчет на {user_email}...')
            
            processor = InvestmentAnalysisProcessor()
            company_name = user_data.get('company_name', 'unknown_company')
            analysis_results = user_data.get('analysis_results')
            qa_history = user_data.get('qa_history', [])
            
            final_report_path = await processor.create_final_report_with_qa(company_name, analysis_results, qa_history)
            
            safe_company_name = processor._sanitize_filename(company_name)
            report_filename = f'investment_analysis_{safe_company_name}_final.docx'
            
            success = await self.email_sender.send_report(user_email, company_name, final_report_path, filename=report_filename)
            
            if os.path.exists(final_report_path):
                os.unlink(final_report_path)
            
            if success:
                await callback_query.message.edit_text(f'✅ Отчет успешно отправлен на {user_email}')
            else:
                await callback_query.message.edit_text(f'❌ Не удалось отправить отчет на {user_email}.\nПопробуйте скачать отчет.')
            
            await callback_query.message.answer('Что бы вы хотели сделать дальше?', reply_markup=FinalActionsKeyboard())
            await UserStates.CHOOSING_FINAL_ACTION.set()
            
        except Exception as e:
            await self.handle_error(callback_query.message, e, "email_sending")

    def register(self, dp: Dispatcher) -> None:
        dp.register_callback_query_handler(
            self.process,
            lambda c: c.data in ['investment_download', 'investment_email', 'investment_back_to_actions'],
            state=UserStates.INVESTMENT_REPORT_OPTIONS,
        )

class FinalActionsHandler(BaseScenario):
    """Обработка финальных действий после получения отчета."""

    async def process(self, callback_query, state, **kwargs) -> None:
        user_id = callback_query.from_user.id
        action = callback_query.data

        await callback_query.answer()

        if action == 'new_company_analysis':
            await callback_query.message.delete()
            await state.finish()
            
            chat_context = ChatContextManager()
            chat_context.end_active_chats(user_id)
            chat_context.cleanup_user_context(user_id)
            
            system_prompts = SystemPrompts()
            system_prompt = system_prompts.get_prompt(SystemPrompt.INVESTMENT)
            chat_context.start_new_chat(user_id, 'investment', system_prompt)
            
            await callback_query.message.answer('Введите название новой компании или опишите ваш запрос для анализа инвестиционной привлекательности.')
            await UserStates.ENTERING_PROMPT.set()
            
        elif action == 'return_to_main_bot':
            await callback_query.message.delete()
            await state.finish()
            
            keyboard = types.InlineKeyboardMarkup()
            keyboard.add(types.InlineKeyboardButton(text="Перейти к основному боту", url="https://t.me/sberallaibot"))
            
            await callback_query.message.answer(
                'Спасибо за использование бота для анализа инвестиционной привлекательности!\n\n'
                'Нажмите кнопку ниже, чтобы перейти к основному боту Сбер CPNB:',
                reply_markup=keyboard
            )

    def register(self, dp: Dispatcher) -> None:
        dp.register_callback_query_handler(
            self.process,
            lambda c: c.data in ['new_company_analysis', 'return_to_main_bot'],
            state=UserStates.CHOOSING_FINAL_ACTION,
        )

# =============================================================================
# MAIN HANDLERS
# =============================================================================

class StartHandler(BaseScenario):
    """Обработка /start команды с проверкой авторизации."""

    async def process(self, message: types.Message, **kwargs) -> None:
        user_id = message.from_user.id
        user_name = f'{message.from_user.first_name} {message.from_user.last_name}'
        logger.info(f'Команда /start от пользователя {user_id} ({user_name})')

        chat_context = ChatContextManager()
        chat_context.end_active_chats(user_id)
        chat_context.cleanup_user_context(user_id)

        # Проверка авторизации
        try:
            is_authorized = await check_user_authorized(user_id)
            
            if not is_authorized:
                logger.info(f'Пользователь {user_id} НЕ найден в базе - перенаправляем на авторизацию')
                await message.answer(
                    "🔒 Доступ к боту ограничен.\n\n"
                    "Для получения доступа пройдите авторизацию в главном боте.",
                    reply_markup=UnauthorizedKeyboard.get_markup()
                )
                return
        
        except Exception as e:
            logger.error(f"Ошибка проверки авторизации для пользователя {user_id}: {e}")
            await message.answer(
                "⚠️ Временные технические неполадки.\n"
                "Попробуйте авторизоваться в главном боте или обратитесь к администратору.",
                reply_markup=UnauthorizedKeyboard.get_markup()
            )
            return
        
        # Если авторизован
        logger.info(f'✅ Пользователь {user_id} авторизован - продолжаем работу')

        await message.answer('Здравствуйте! Добро пожаловать в бот для анализа инвестиционной привлекательности.\n\nВведите название компании или опишите ваш запрос для анализа.')
        
        # Автоматически устанавливаем тему как investment
        system_prompts = SystemPrompts()
        system_prompt = system_prompts.get_prompt(SystemPrompt.INVESTMENT)
        chat_context.start_new_chat(user_id, 'investment', system_prompt)
        
        await UserStates.ENTERING_PROMPT.set()

    def register(self, dp: Dispatcher) -> None:
        dp.register_message_handler(self.process, commands=['start'], state='*')

class ProcessingEnterPromptHandler(BaseScenario):
    """Обработка ввода текстового промпта пользователем."""

    async def process(self, message: types.Message, state: FSMContext, **kwargs) -> None:
        user_id = message.from_user.id
        
        # Автоматически устанавливаем investment тему
        await state.update_data(chosen_topic='investment', chosen_model='chatgpt', user_query=message.text)

        logger.info(f'Получен текстовый запрос от {user_id}: тема=investment')

        file_message = await message.answer('Хотите ли вы прикрепить файл (PDF, Word, PPT) для анализа?', reply_markup=FileAttachKeyboard())
        await state.update_data(file_message_id=file_message.message_id)
        await UserStates.ATTACHING_FILE.set()

    def register(self, dp: Dispatcher) -> None:
        dp.register_message_handler(self.process, content_types=['text'], state=UserStates.ENTERING_PROMPT)

class AttachFileHandler(BaseScenario):
    """Обработчик прикрепления файла."""

    async def process(self, callback_query: types.CallbackQuery, state: FSMContext, **kwargs) -> None:
        user_id = callback_query.from_user.id
        user_data = await state.get_data()
        await callback_query.answer()
        await self.delete_message_by_id(user_id, user_data.get('file_message_id'))
        
        if callback_query.data == 'attach_file':
            file_prompt = await callback_query.message.answer('Пожалуйста, загрузите файл (PDF, Word, PPT):')
            await state.update_data(file_prompt_id=file_prompt.message_id)
            await UserStates.UPLOADING_FILE.set()
        else:
            # Без файла
            await self.process_investment_analysis(callback_query.message, state, file_content='')

    def register(self, dp: Dispatcher) -> None:
        dp.register_callback_query_handler(
            self.process,
            lambda c: c.data in ['attach_file', 'no_file'],
            state=[UserStates.ATTACHING_FILE, UserStates.ATTACHING_FILE_CONTINUE],
        )

class UploadFileHandler(BaseScenario):
    """Обработчик загрузки файла."""

    async def process(self, message: types.Message, state: FSMContext, **kwargs) -> None:
        user_id = message.from_user.id
        user_data = await state.get_data()
        await self.delete_message_by_id(user_id, user_data.get('file_prompt_id'))
        
        if not message.document:
            await message.answer('Пожалуйста, загрузите файл в формате PDF, Word или PowerPoint.')
            return
            
        try:
            processing_msg = await message.answer('Идет обработка файла...')
            file_content = await FileProcessor.extract_text_from_file(message.document, self.bot)
            await state.update_data(processing_msg_id=processing_msg.message_id)
            await self.process_investment_analysis(message, state, file_content)
        except ValueError as e:
            logger.error(f'Ошибка обработки файла: {e}')
            await message.answer('Произошла ошибка при обработке файла. Продолжите использование нажав команду /start')
            await self.bot.send_message(chat_id=config.OWNER_ID, text=str(e))

    def register(self, dp: Dispatcher) -> None:
        dp.register_message_handler(self.process, content_types=['document'], state=[UserStates.UPLOADING_FILE, UserStates.UPLOADING_FILE_CONTINUE])

class ResetStateHandler(BaseScenario):
    """Обработка команды /reset."""

    async def process(self, message: types.Message, state: FSMContext, **kwargs) -> None:
        user_id = message.from_user.id
        logger.info(f'Пользователь {user_id} запросил сброс состояния')

        await state.finish()
        await message.answer('Состояние сброшено. Введите название компании или опишите ваш запрос для анализа инвестиционной привлекательности.')
        await UserStates.ENTERING_PROMPT.set()

    def register(self, dp: Dispatcher) -> None:
        dp.register_message_handler(self.process, commands=['reset'], state='*')

# =============================================================================
# BOT MANAGER
# =============================================================================

class BotManager:
    scenarios: Dict[str, BaseScenario] = {}

    def __init__(self, bot: Bot, dp: Dispatcher) -> None:
        self.bot = bot
        self.dp = dp

        # Основные сценарии
        main_scenarios = {
            'start': StartHandler,
            'enter_prompt': ProcessingEnterPromptHandler,
            'attach_file': AttachFileHandler,
            'upload_file': UploadFileHandler,
            'reset_state': ResetStateHandler,
        }

        # Инвестиционные сценарии
        investment_scenarios = {
            'investment_actions': InvestmentActionsHandler,    
            'investment_qa': InvestmentQAHandler, 
            'back_to_investment_actions': BackToInvestmentActionsHandler, 
            'investment_report': InvestmentReportHandler, 
            'final_actions': FinalActionsHandler,
        }

        self._setup_middlewares()

        # Регистрируем все сценарии
        all_scenarios = [
            ('main', main_scenarios),
            ('investment', investment_scenarios),
        ]

        for scenario_group, scenarios in all_scenarios:
            for scenario_name, scenario_class in scenarios.items():
                full_name = f'{scenario_group}_{scenario_name}'
                logger.info(f'Registering scenario: {full_name}')
                scenario_instance = scenario_class(bot)
                self._register_scenario(full_name, scenario_instance)

        # Регистрируем все обработчики
        logger.info("Starting registration of all handlers...")
        for scenario_name, scenario in self.scenarios.items():
            logger.info(f'Calling register() for scenario: {scenario_name}')
            try:
                scenario.register(dp)
                logger.info(f'Successfully registered: {scenario_name}')
            except Exception as e:
                logger.error(f'Failed to register {scenario_name}: {e}')

        logger.info(f"Total scenarios registered: {len(self.scenarios)}")

    def _register_scenario(self, name: str, scenario: BaseScenario) -> None:
        self.scenarios[name] = scenario
        logger.info(f'Scenario {name} added to scenarios dict')

    def _setup_middlewares(self) -> None:
        self.dp.middleware.setup(AccessMiddleware())

# =============================================================================
# MAIN
# =============================================================================

if __name__ == '__main__':
    from aiogram import executor
    
    async def on_startup(dp):
        """Инициализация при запуске бота"""
        try:
            config = Config()
            sql_connection = config.SQL_CONNECTION_STRING_READER
            await init_auth_system(sql_connection)
            logger.info("✅ Система авторизации инициализирована")
        except Exception as e:
            logger.error(f"❌ Ошибка инициализации авторизации: {e}")

    config = Config()
    bot = Bot(token=config.TOKEN)
    dp = Dispatcher(bot, storage=MemoryStorage())

    BotManager(bot, dp)

    # Запуск с инициализацией авторизации
    executor.start_polling(dp, skip_updates=True, on_startup=on_startup)